"""
Generate a DOCX internship practice report for the LunchTogether application.

Formatting follows university standards:
  - A4, margins: left 30mm, right 10mm, top/bottom 20mm
  - Times New Roman 14pt, single line spacing
  - Page numbers bottom-right, starting from page 2
  - Sequential numbering for figures and tables
  - Syntax-highlighted code listings

Usage:
    pip install python-docx matplotlib
    python generate_report.py

Output: practice_report.docx in the project root.
"""

import re
from pathlib import Path

import matplotlib
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

matplotlib.use("Agg")

PROJECT_ROOT = Path(__file__).parent
DIAGRAMS_DIR = PROJECT_ROOT / "_report_diagrams"
OUTPUT_FILE = PROJECT_ROOT / "practice_report.docx"

# ---------------------------------------------------------------------------
# Sequential numbering counters
# ---------------------------------------------------------------------------
_fig_counter = 0
_table_counter = 0
_listing_section_counters: dict[str, int] = {}


def _next_fig():
    global _fig_counter
    _fig_counter += 1
    return _fig_counter


def _next_table():
    global _table_counter
    _table_counter += 1
    return _table_counter


def _next_listing(section: str) -> str:
    _listing_section_counters[section] = _listing_section_counters.get(section, 0) + 1
    return f"{section}.{_listing_section_counters[section]}"


# ---------------------------------------------------------------------------
# Syntax highlighting
# ---------------------------------------------------------------------------

COLOR_KEYWORD = RGBColor(0, 0, 180)
COLOR_STRING = RGBColor(0, 128, 0)
COLOR_COMMENT = RGBColor(128, 128, 128)
COLOR_DECORATOR = RGBColor(128, 0, 128)
COLOR_NUMBER = RGBColor(128, 64, 0)
COLOR_PLAIN = RGBColor(0, 0, 0)

PYTHON_KEYWORDS = {
    "False", "None", "True", "and", "as", "assert", "async", "await",
    "break", "class", "continue", "def", "del", "elif", "else", "except",
    "finally", "for", "from", "global", "if", "import", "in", "is",
    "lambda", "nonlocal", "not", "or", "pass", "raise", "return",
    "try", "while", "with", "yield", "self",
}

TS_KEYWORDS = {
    "as", "async", "await", "break", "case", "catch", "class", "const",
    "continue", "debugger", "default", "delete", "do", "else", "enum",
    "export", "extends", "false", "finally", "for", "from", "function",
    "if", "implements", "import", "in", "instanceof", "interface", "keyof",
    "let", "module", "namespace", "new", "null", "of", "private",
    "protected", "public", "readonly", "return", "static", "super",
    "switch", "this", "throw", "true", "try", "type", "typeof",
    "undefined", "var", "void", "while", "with", "yield",
}

_TOKEN_RE = re.compile(
    r'("""[\s\S]*?"""|'
    r"'''[\s\S]*?''')|"        # 1: triple-quoted strings
    r'("(?:[^"\\]|\\.)*"|'
    r"'(?:[^'\\]|\\.)*'|"
    r"`(?:[^`\\]|\\.)*`)|"     # 2: quoted strings
    r"(#.*)|"                  # 3: Python comment
    r"(//.*)|"                 # 4: JS/TS comment
    r"(@\w[\w.]*)|"            # 5: decorator
    r"(\b\d+\.?\d*\b)|"        # 6: number
    r"(\b[A-Za-z_]\w*\b)|"     # 7: identifier
    r"(\s+)|"                  # 8: whitespace
    r"(.)"                     # 9: other
)


def tokenize_line(line, lang="python"):
    keywords = PYTHON_KEYWORDS if lang == "python" else TS_KEYWORDS
    tokens = []
    for m in _TOKEN_RE.finditer(line):
        text = m.group()
        if m.group(1) or m.group(2):
            tokens.append((text, "string"))
        elif m.group(3) or m.group(4):
            tokens.append((text, "comment"))
        elif m.group(5):
            tokens.append((text, "decorator"))
        elif m.group(6):
            tokens.append((text, "number"))
        elif m.group(7):
            if text in keywords:
                tokens.append((text, "keyword"))
            else:
                tokens.append((text, "plain"))
        else:
            tokens.append((text, "plain"))
    return tokens


def _apply_token_style(run, token_type):
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    if token_type == "keyword":
        run.font.color.rgb = COLOR_KEYWORD
        run.bold = True
    elif token_type == "string":
        run.font.color.rgb = COLOR_STRING
    elif token_type == "comment":
        run.font.color.rgb = COLOR_COMMENT
        run.italic = True
    elif token_type == "decorator":
        run.font.color.rgb = COLOR_DECORATOR
    elif token_type == "number":
        run.font.color.rgb = COLOR_NUMBER
    else:
        run.font.color.rgb = COLOR_PLAIN


# ---------------------------------------------------------------------------
# Document formatting helpers
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="CCCCCC", style="single", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def add_body_text(doc, text, bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    return para


def add_chapter_heading(doc, text):
    """РОЗДІЛ / ВСТУП / ВИСНОВКИ — centered, bold, no period, page break."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.page_break_before = True
    pf.first_line_indent = Cm(0)
    run = para.add_run(text.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    para.style = doc.styles["Heading 1"]
    for r in para.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_section_heading(doc, text):
    """1.1. Title. — indented, bold, WITH period at end."""
    display = text if text.endswith(".") else text + "."
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(1.25)
    run = para.add_run(display)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    para.style = doc.styles["Heading 2"]
    for r in para.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_listing_caption(doc, section, description):
    """Listing caption ABOVE the code block, centered, section-relative number."""
    num = _next_listing(section)
    label = f"Лістинг {num}. {description}"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    return para


def add_code_annotation(doc, text):
    """Brief annotation before a code listing."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_code_listing(doc, code_text, lang="python"):
    """Add a syntax-highlighted code block inside a bordered table cell."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FAFAFA")
    _set_cell_borders(cell, color="CCCCCC")

    cell.paragraphs[0].clear()
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)

        if not line.strip():
            run = para.add_run(" ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        tokens = tokenize_line(line, lang)
        for text, token_type in tokens:
            run = para.add_run(text)
            _apply_token_style(run, token_type)

    # Small gap after code
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return table


def add_figure_caption(doc, description):
    """Figure caption BELOW the image, centered, sequential number."""
    num = _next_fig()
    label = f"Рис. {num}. {description}"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    return para


def add_table_caption(doc, description):
    """Table caption ABOVE the table, right-aligned, sequential number."""
    num = _next_table()
    label = f"Таблиця {num}. {description}"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    return para


def add_screenshot_placeholder(doc, description, height_cm=6):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_borders(cell, color="999999", style="dashed", sz="6")

    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("[Місце для зображення]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(153, 153, 153)

    add_figure_caption(doc, description)


def add_image(doc, image_path, description, width_cm=15):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_figure_caption(doc, description)


def add_toc(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("ЗМІСТ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    para2 = doc.add_paragraph()
    para2.paragraph_format.first_line_indent = Cm(0)

    for field_type in ("begin", "separate", "end"):
        run = para2.add_run()
        if field_type == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            run._r.append(fc)
            run2 = para2.add_run()
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = ' TOC \\o "1-2" \\h \\z \\u '
            run2._r.append(it)
        elif field_type == "separate":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "separate")
            run._r.append(fc)
            hint = para2.add_run("(Оновіть зміст: ПКМ → Оновити поле)")
            hint.font.name = "Times New Roman"
            hint.font.size = Pt(12)
            hint.font.color.rgb = RGBColor(128, 128, 128)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            run._r.append(fc)


def add_page_number_footer(section):
    """Add a right-aligned page number to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.first_line_indent = Cm(0)

    run = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fc_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)

    run3 = para.add_run()
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fc_end)

    for r in para.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def make_data_table(doc, headers, rows, col_widths=None):
    """Create a formatted data table. Caption should be added BEFORE calling this."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
        _set_cell_shading(cell, "D9E2F3")

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    return table


def read_code_file(relative_path, start=None, end=None):
    full_path = PROJECT_ROOT / relative_path
    if not full_path.exists():
        return f"# File not found: {relative_path}"
    text = full_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    if start is not None or end is not None:
        lines = lines[(start or 1) - 1 : end]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Diagram generation (matplotlib)
# ---------------------------------------------------------------------------


def generate_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")

    boxes = [
        (1, 5.5, 8, 1.0, "Клієнт (React + TypeScript)\nRedux Toolkit / RTK Query / React Router / shadcn/ui", "#E3F2FD"),
        (1, 3.8, 8, 1.0, "REST API (FastAPI + Python)\nPydantic / JWT / Middleware / CORS", "#E8F5E9"),
        (1, 2.1, 3.5, 1.0, "Бізнес-логіка\nWorkflow-класи", "#FFF3E0"),
        (5.5, 2.1, 3.5, 1.0, "Доступ до даних\nRepository-класи", "#FFF3E0"),
        (1, 0.4, 8, 1.0, "База даних (PostgreSQL)\nSQLAlchemy ORM / Alembic міграції", "#FCE4EC"),
    ]

    for x, y, w, h, text, color in boxes:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.1",
            facecolor=color, edgecolor="#333333", linewidth=1.5,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=10, fontfamily="sans-serif")

    arr = dict(arrowstyle="->", color="#555555", lw=2)
    ax.annotate("", xy=(5, 5.5), xytext=(5, 4.8), arrowprops=arr)
    ax.annotate("", xy=(2.75, 3.8), xytext=(2.75, 3.1), arrowprops=arr)
    ax.annotate("", xy=(7.25, 3.8), xytext=(7.25, 3.1), arrowprops=arr)
    ax.annotate("", xy=(2.75, 2.1), xytext=(2.75, 1.4), arrowprops=arr)
    ax.annotate("", xy=(7.25, 2.1), xytext=(7.25, 1.4), arrowprops=arr)

    path = DIAGRAMS_DIR / "architecture.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_er_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis("off")

    entities = [
        (0.5, 6.5, 2.5, 2.0, "Users\n─────────\nid (PK)\nemail\nfull_name\nrole\nis_active"),
        (4.0, 7.0, 2.5, 1.5, "Groups\n─────────\nid (PK)\nname\nowner_id (FK)"),
        (7.5, 7.0, 3.0, 1.5, "GroupMembers\n─────────\nid (PK)\nuser_id (FK)\ngroup_id (FK)"),
        (7.5, 5.0, 3.0, 1.5, "Permissions\n─────────\ngroup_member_id\npermission_type\nlevel"),
        (0.5, 3.5, 2.5, 1.5, "Restaurants\n─────────\nid (PK)\nname\ngroup_id (FK)"),
        (4.0, 3.5, 2.5, 1.8, "Orders\n─────────\nid (PK)\ngroup_id (FK)\ninitiator_id (FK)\nstatus"),
        (7.5, 3.0, 3.0, 1.5, "OrderItems\n─────────\norder_id (FK)\nuser_id (FK)\nname, price"),
        (0.5, 1.0, 2.5, 1.5, "Dishes\n─────────\nid (PK)\nname, price\nrestaurant_id (FK)"),
        (4.0, 1.0, 2.5, 1.5, "Balances\n─────────\nuser_id (FK)\ngroup_id (FK)\namount"),
        (7.5, 1.0, 3.0, 1.5, "BalanceHistory\n─────────\nbalance_id (FK)\namount\nchange_type"),
    ]

    for x, y, w, h, text in entities:
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.08",
            facecolor="#FFFDE7", edgecolor="#333333", linewidth=1.2,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=7.5, fontfamily="monospace")

    lines = [
        ((3.0, 7.5), (4.0, 7.5)),
        ((6.5, 7.5), (7.5, 7.5)),
        ((9.0, 7.0), (9.0, 6.5)),
        ((5.25, 7.0), (5.25, 5.3)),
        ((1.75, 6.5), (1.75, 5.0)),
        ((6.5, 4.5), (7.5, 4.0)),
        ((5.25, 3.5), (5.25, 2.5)),
        ((6.5, 1.5), (7.5, 1.5)),
        ((1.75, 3.5), (1.75, 2.5)),
    ]
    for (x1, y1), (x2, y2) in lines:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                     arrowprops=dict(arrowstyle="-", color="#666666", lw=1.2))

    path = DIAGRAMS_DIR / "er_diagram.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_permission_flowchart():
    fig, ax = plt.subplots(1, 1, figsize=(8, 10))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 11)
    ax.axis("off")

    def box(x, y, w, h, text, color="#E3F2FD"):
        r = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.1",
            facecolor=color, edgecolor="#333", linewidth=1.2,
        )
        ax.add_patch(r)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=9, fontfamily="sans-serif")

    def diamond(cx, cy, text, color="#FFF9C4"):
        d = plt.Polygon(
            [(cx, cy + 0.6), (cx + 1.5, cy), (cx, cy - 0.6), (cx - 1.5, cy)],
            facecolor=color, edgecolor="#333", linewidth=1.2,
        )
        ax.add_patch(d)
        ax.text(cx, cy, text, ha="center", va="center", fontsize=8, fontfamily="sans-serif")

    arr = dict(arrowstyle="->", color="#555", lw=1.5)

    box(2.5, 9.8, 3, 0.7, "Запит користувача", "#BBDEFB")
    diamond(4, 8.5, "Користувач\nAdmin?")
    box(6.0, 8.2, 1.8, 0.6, "ДОЗВОЛЕНО", "#C8E6C9")
    diamond(4, 6.8, "Є учасником\nгрупи?")
    box(6.0, 6.5, 1.8, 0.6, "ЗАБОРОНЕНО", "#FFCDD2")
    diamond(4, 5.1, "Має потрібний\nPermissionType?")
    box(6.0, 4.8, 1.8, 0.6, "ЗАБОРОНЕНО", "#FFCDD2")
    diamond(4, 3.4, "Рівень дозволу\nдостатній?")
    box(6.0, 3.1, 1.8, 0.6, "ЗАБОРОНЕНО", "#FFCDD2")
    box(2.5, 1.8, 3, 0.7, "ДОЗВОЛЕНО", "#C8E6C9")

    ax.annotate("", xy=(4, 9.1), xytext=(4, 9.8), arrowprops=arr)
    ax.annotate("", xy=(6.0, 8.5), xytext=(5.5, 8.5), arrowprops=arr)
    ax.text(5.6, 8.7, "Так", fontsize=8)
    ax.annotate("", xy=(4, 7.4), xytext=(4, 7.9), arrowprops=arr)
    ax.text(3.5, 7.7, "Ні", fontsize=8)
    ax.annotate("", xy=(6.0, 6.8), xytext=(5.5, 6.8), arrowprops=arr)
    ax.text(5.6, 7.0, "Ні", fontsize=8)
    ax.annotate("", xy=(4, 5.7), xytext=(4, 6.2), arrowprops=arr)
    ax.text(3.5, 6.0, "Так", fontsize=8)
    ax.annotate("", xy=(6.0, 5.1), xytext=(5.5, 5.1), arrowprops=arr)
    ax.text(5.6, 5.3, "Ні", fontsize=8)
    ax.annotate("", xy=(4, 4.0), xytext=(4, 4.5), arrowprops=arr)
    ax.text(3.5, 4.3, "Так", fontsize=8)
    ax.annotate("", xy=(6.0, 3.4), xytext=(5.5, 3.4), arrowprops=arr)
    ax.text(5.6, 3.6, "Ні", fontsize=8)
    ax.annotate("", xy=(4, 2.5), xytext=(4, 2.8), arrowprops=arr)
    ax.text(3.5, 2.7, "Так", fontsize=8)

    path = DIAGRAMS_DIR / "permission_flow.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_order_lifecycle():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")

    states = [
        (0.5, 2.0, "Initiated", "#BBDEFB"),
        (3.0, 2.0, "Confirmed", "#FFF9C4"),
        (5.5, 2.0, "Ordered", "#E1BEE7"),
        (8.0, 2.0, "Finished", "#C8E6C9"),
        (5.5, 0.3, "Cancelled", "#FFCDD2"),
    ]

    for x, y, label, color in states:
        rect = mpatches.FancyBboxPatch(
            (x, y), 2.0, 0.9, boxstyle="round,pad=0.12",
            facecolor=color, edgecolor="#333", linewidth=1.5,
        )
        ax.add_patch(rect)
        ax.text(x + 1.0, y + 0.45, label, ha="center", va="center",
                fontsize=11, fontweight="bold", fontfamily="sans-serif")

    arr = dict(arrowstyle="-|>", color="#333", lw=2)
    arr_cancel = dict(arrowstyle="-|>", color="#D32F2F", lw=1.5, linestyle="dashed")

    ax.annotate("", xy=(3.0, 2.45), xytext=(2.5, 2.45), arrowprops=arr)
    ax.annotate("", xy=(5.5, 2.45), xytext=(5.0, 2.45), arrowprops=arr)
    ax.annotate("", xy=(8.0, 2.45), xytext=(7.5, 2.45), arrowprops=arr)
    ax.annotate("", xy=(6.5, 1.2), xytext=(1.5, 2.0), arrowprops=arr_cancel)
    ax.annotate("", xy=(6.5, 1.2), xytext=(4.0, 2.0), arrowprops=arr_cancel)
    ax.annotate("", xy=(6.5, 1.2), xytext=(6.5, 2.0), arrowprops=arr_cancel)

    ax.text(1.5, 3.3, "confirm", fontsize=9, ha="center", color="#555", fontstyle="italic")
    ax.text(4.0, 3.3, "order", fontsize=9, ha="center", color="#555", fontstyle="italic")
    ax.text(6.5, 3.3, "finish", fontsize=9, ha="center", color="#555", fontstyle="italic")
    ax.text(3.0, 0.7, "cancel", fontsize=9, ha="center", color="#D32F2F", fontstyle="italic")

    path = DIAGRAMS_DIR / "order_lifecycle.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ---------------------------------------------------------------------------
# Main report builder
# ---------------------------------------------------------------------------


def build_report():
    DIAGRAMS_DIR.mkdir(exist_ok=True)

    arch_path = generate_architecture_diagram()
    er_path = generate_er_diagram()
    perm_path = generate_permission_flowchart()
    order_path = generate_order_lifecycle()

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # "Different first page" — title page has no footer
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)

    add_page_number_footer(section)

    # ── Styles ──────────────────────────────────────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(14)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for lvl in ("Heading 1", "Heading 2"):
        s = doc.styles[lvl]
        s.font.name = "Times New Roman"
        s.font.size = Pt(14)
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.font.bold = True
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ===========================================================
    #  TITLE PAGE
    # ===========================================================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("ЗВІТ\nпро проходження навчальної практики")
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(
        'Тема: "Розробка веб-додатку LunchTogether\n'
        'для організації спільних обідів у колективах"'
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(
        "Виконав: студент групи ____\n"
        "____________________\n\n"
        "Керівник практики:\n"
        "____________________"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    # ===========================================================
    #  TABLE OF CONTENTS
    # ===========================================================
    doc.add_page_break()
    add_toc(doc)

    # ===========================================================
    #  ВСТУП
    # ===========================================================
    add_chapter_heading(doc, "ВСТУП")

    add_body_text(doc,
        "Метою навчальної практики є набуття практичних навичок розробки сучасних веб-додатків "
        "повного циклу із застосуванням актуальних технологій та архітектурних підходів. "
        "Завданням практики є проєктування та реалізація веб-додатку «LunchTogether» — системи "
        "для організації спільних обідів у робочих колективах."
    )
    add_body_text(doc,
        "Основні завдання практики включають: аналіз предметної області та існуючих рішень; "
        "формування функціональних та нефункціональних вимог; проєктування архітектури системи "
        "та бази даних; розробку серверної частини на базі Python та FastAPI; розробку клієнтської "
        "частини на базі React та TypeScript; реалізацію системи автентифікації, авторизації "
        "та гнучких групових дозволів."
    )
    add_body_text(doc,
        "Базою практики є кафедра інформаційних технологій університету. Під час практики "
        "було виконано повний цикл розробки програмного забезпечення — від аналізу вимог "
        "до реалізації готового продукту."
    )
    add_body_text(doc,
        "Звіт складається з чотирьох розділів. У першому розділі проведено аналіз предметної "
        "області та сформовано вимоги до системи. У другому розділі описано проєктування "
        "архітектури. Третій розділ присвячено розробці серверної частини. У четвертому розділі "
        "розглядається розробка клієнтської частини додатку."
    )

    # ===========================================================
    #  РОЗДІЛ 1
    # ===========================================================
    add_chapter_heading(doc, "РОЗДІЛ 1. АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ ТА ПОСТАНОВКА ЗАДАЧІ")

    add_section_heading(doc, "1.1. Проблематика організації спільних обідів у колективах")
    add_body_text(doc,
        "Організація спільних обідів у робочих колективах є поширеною практикою, що сприяє "
        "зміцненню командного духу та підвищенню ефективності комунікації. Проте цей процес "
        "часто супроводжується низкою організаційних проблем."
    )
    add_body_text(doc,
        "По-перше, виникає проблема координації замовлень: збір інформації про бажані страви "
        "від кожного учасника є трудомістким процесом, що зазвичай відбувається через месенджери "
        "або усне опитування. По-друге, розподіл витрат між учасниками потребує окремого обліку, "
        "оскільки вартість доставки та різних страв може суттєво відрізнятися. По-третє, "
        "відсутність єдиної платформи призводить до втрати інформації про попередні замовлення "
        "та вподобання учасників."
    )
    add_body_text(doc,
        "Додатковою проблемою є управління правами доступу: не кожен учасник групи повинен "
        "мати можливість ініціювати замовлення або змінювати баланси інших учасників. Відсутність "
        "гнучкої системи дозволів обмежує можливості делегування обов'язків."
    )

    add_section_heading(doc, "1.2. Аналіз існуючих рішень")
    add_body_text(doc,
        "На ринку існує кілька підходів до розв'язання описаних проблем. Месенджери (Telegram, "
        "Slack) дозволяють координувати замовлення через групові чати, але не забезпечують "
        "структурованого обліку та автоматизації. Спеціалізовані сервіси доставки їжі (Glovo, "
        "Bolt Food) орієнтовані на індивідуальні замовлення і не підтримують групову координацію."
    )
    add_body_text(doc,
        "Таблиці Google Sheets можуть бути адаптовані для ведення обліку, але потребують значних "
        "зусиль для налаштування та не забезпечують зручного інтерфейсу. Жодне з існуючих "
        "рішень не поєднує всі необхідні функції: групове управління, систему дозволів, "
        "життєвий цикл замовлень та автоматичний облік балансів."
    )

    add_section_heading(doc, "1.3. Функціональні вимоги до системи")
    add_body_text(doc,
        "На основі аналізу предметної області було сформовано перелік функціональних вимог, "
        "що структуровані за модулями (епіками):"
    )

    add_table_caption(doc, "Функціональні вимоги до системи")
    make_data_table(doc,
        ["№", "Модуль", "Опис функціональності"],
        [
            ("1", "Автентифікація", "Реєстрація, вхід/вихід через JWT, збереження сесії"),
            ("2", "Система дозволів", "Ролі (Admin/User), групові дозволи (Members, Orders, Balances, Analytics, Restaurants)"),
            ("3", "Управління користувачами", "Адміністрування користувачів, зміна ролей"),
            ("4", "Групи", "Створення груп (до 5), управління учасниками (до 25), запрошення"),
            ("5", "Ресторани", "CRUD ресторанів та страв в межах групи"),
            ("6", "Баланси", "Облік балансів, ручне коригування, історія змін"),
            ("7", "Замовлення", "Життєвий цикл замовлення, управління стравами, розподіл доставки"),
            ("8", "Дашборд групи", "Аналітика групи, активне замовлення, баланс"),
            ("9", "Дашборд користувача", "Загальна аналітика, налаштування, навігація"),
        ],
    )

    add_section_heading(doc, "1.4. Нефункціональні вимоги та обмеження")
    add_body_text(doc,
        "До нефункціональних вимог системи відносяться: забезпечення безпеки даних через "
        "JWT-автентифікацію з HTTP-only cookies; підтримка асинхронної обробки запитів; "
        "валідація вхідних даних на рівні API через Pydantic-схеми; типобезпека на клієнтській "
        "частині через TypeScript; адаптивний інтерфейс."
    )
    add_body_text(doc,
        "Технологічні обмеження включають: PostgreSQL як СУБД, Python 3.12+ для серверної "
        "частини, React 19 з TypeScript для клієнтської частини. Валюта — українська гривня (₴)."
    )

    # ===========================================================
    #  РОЗДІЛ 2
    # ===========================================================
    add_chapter_heading(doc, "РОЗДІЛ 2. ПРОЄКТУВАННЯ АРХІТЕКТУРИ СИСТЕМИ")

    add_section_heading(doc, "2.1. Загальна архітектура та структура проєкту")
    add_body_text(doc,
        "Система LunchTogether побудована за клієнт-серверною архітектурою з чітким "
        "розділенням відповідальності між шарами. Клієнтська частина (React SPA) взаємодіє "
        "з серверною частиною (FastAPI REST API) через HTTP-запити. Серверна частина "
        "використовує шаблон Repository для доступу до даних та Workflow-класи для "
        "інкапсуляції бізнес-логіки."
    )

    add_image(doc, arch_path, "Загальна архітектура системи LunchTogether", width_cm=14)

    add_body_text(doc,
        "Проєкт організовано як монорепозиторій з двома основними директоріями: backend/ та "
        "frontend/. Серверна частина структурована за шарами: models, schemas, repositories, "
        "workflows та api. Клієнтська частина організована за модулями функціональності: "
        "auth, group, order, restaurant, balance, dashboard."
    )

    add_section_heading(doc, "2.2. Проєктування бази даних")
    add_body_text(doc,
        "База даних PostgreSQL містить 10 основних таблиць, пов'язаних зовнішніми ключами. "
        "Усі таблиці використовують UUID як первинний ключ та мають поля created_at і "
        "updated_at для аудиту. Міграції керуються за допомогою Alembic."
    )

    add_image(doc, er_path, "ER-діаграма бази даних", width_cm=15)

    add_table_caption(doc, "Основні таблиці бази даних")
    make_data_table(doc,
        ["Таблиця", "Призначення", "Ключові поля"],
        [
            ("users", "Користувачі системи", "email, role, full_name"),
            ("groups", "Групи для спільних обідів", "name, owner_id"),
            ("group_members", "Учасники груп", "user_id, group_id"),
            ("group_member_permissions", "Дозволи учасників", "permission_type, level"),
            ("group_invitations", "Запрошення до груп", "invitee_email, token, status"),
            ("restaurants", "Ресторани в межах групи", "name, group_id"),
            ("dishes", "Страви ресторану", "name, price, restaurant_id"),
            ("orders", "Замовлення", "group_id, status, initiator_id"),
            ("order_items", "Позиції замовлення", "order_id, user_id, name, price"),
            ("balances", "Баланси користувачів", "user_id, group_id, amount"),
        ],
    )

    add_section_heading(doc, "2.3. Проєктування REST API")
    add_body_text(doc,
        "REST API побудовано на основі FastAPI з використанням APIRouter для модульної "
        "організації ендпоінтів. Усі ендпоінти повертають JSON-відповіді з Pydantic-схемами "
        "для автоматичної валідації та документації через OpenAPI (Swagger)."
    )

    add_table_caption(doc, "Основні ендпоінти REST API")
    make_data_table(doc,
        ["Метод", "Шлях", "Опис"],
        [
            ("POST", "/api/auth/register", "Реєстрація користувача"),
            ("POST", "/api/auth/login", "Вхід в систему"),
            ("GET", "/api/groups", "Список груп користувача"),
            ("POST", "/api/groups", "Створення групи"),
            ("GET", "/api/groups/{id}/orders", "Список замовлень групи"),
            ("POST", "/api/groups/{id}/orders", "Створення замовлення"),
            ("POST", "/api/groups/{id}/orders/{oid}/status", "Зміна статусу замовлення"),
            ("GET", "/api/groups/{id}/balances", "Баланси учасників групи"),
        ],
    )

    add_section_heading(doc, "2.4. Система прав та дозволів")
    add_body_text(doc,
        "Система прав реалізована на двох рівнях. На рівні додатку визначено ролі Admin та User. "
        "На рівні групи дозволи зберігаються в таблиці group_member_permissions з п'ятьма типами: "
        "Members, Orders, Balances, Analytics, Restaurants. Кожен тип має свої рівні доступу "
        "(Editor, Viewer, Initiator, Participant, None)."
    )

    add_image(doc, perm_path, "Алгоритм перевірки прав доступу", width_cm=10)

    add_body_text(doc,
        "Для спрощення призначення дозволів реалізовано систему пресетів ролей: Admin (повний "
        "доступ), SupervisorMember (перегляд та ініціювання), Member (базова участь)."
    )

    # ===========================================================
    #  РОЗДІЛ 3
    # ===========================================================
    add_chapter_heading(doc, "РОЗДІЛ 3. РОЗРОБКА СЕРВЕРНОЇ ЧАСТИНИ")

    add_section_heading(doc, "3.1. Налаштування FastAPI та структура бекенду")
    add_body_text(doc,
        "Серверна частина побудована на фреймворку FastAPI — сучасному асинхронному "
        "Python-фреймворку для REST API. FastAPI забезпечує автоматичну генерацію "
        "OpenAPI-документації, валідацію запитів через Pydantic та підтримку async/await."
    )

    add_listing_caption(doc, "3.1", "Точка входу FastAPI-додатку (main.py)")
    add_code_annotation(doc,
        "Головний модуль серверного додатку, що ініціалізує FastAPI із фабричною функцією "
        "create_app(), налаштовує CORS, middleware логування та обробки помилок, "
        "підключає маршрутизатор API та інтегрує Sentry для моніторингу."
    )
    add_code_listing(doc, read_code_file("backend/app/main.py"), "python")

    add_section_heading(doc, "3.2. Моделі даних SQLAlchemy")
    add_body_text(doc,
        "Моделі даних визначено з використанням декларативного стилю SQLAlchemy 2.0 "
        "з типізованими атрибутами (Mapped, mapped_column). Усі моделі наслідують BaseModel, "
        "що забезпечує спільні поля id (UUID), created_at та updated_at."
    )

    add_listing_caption(doc, "3.2", "Модель користувача (User)")
    add_code_annotation(doc,
        "Модель User визначає поля email, full_name, role (Admin/User), is_active, is_verified. "
        "Властивість is_admin забезпечує зворотну сумісність. Зв'язки: owned_groups, group_memberships."
    )
    add_code_listing(doc, read_code_file("backend/app/models/user.py"), "python")

    add_listing_caption(doc, "3.2", "Модель замовлення (Order та OrderItem)")
    add_code_annotation(doc,
        "Модель Order зберігає інформацію про замовлення: група, ресторан, ініціатор, статус "
        "та вартість доставки. OrderItem описує окрему позицію: назва, ціна, кількість, страва."
    )
    add_code_listing(doc, read_code_file("backend/app/models/order.py", end=65), "python")

    add_listing_caption(doc, "3.2", "Моделі GroupMember та GroupMemberPermission")
    add_code_annotation(doc,
        "GroupMember пов'язує користувача з групою. GroupMemberPermission зберігає окремі "
        "дозволи (permission_type, level) з унікальним ключем по group_member_id та permission_type. "
        "Метод get_permission() шукає рівень доступу для заданого типу дозволу."
    )
    add_code_listing(doc, read_code_file("backend/app/models/group.py", start=65, end=130), "python")

    add_section_heading(doc, "3.3. Шаблон Repository для доступу до даних")
    add_body_text(doc,
        "Для доступу до даних реалізовано шаблон Repository, який інкапсулює всі операції "
        "з базою даних. Базовий репозиторій BaseRepository є узагальненим (Generic) класом, "
        "що надає CRUD-операції для будь-якої моделі."
    )

    add_listing_caption(doc, "3.3", "Базовий репозиторій (BaseRepository)")
    add_code_annotation(doc,
        "Узагальнений клас BaseRepository[ModelType] реалізує стандартні операції: get_by_id, "
        "get, get_multi (з пагінацією), create, update, delete. Використовує AsyncSession "
        "для асинхронного доступу до PostgreSQL."
    )
    add_code_listing(doc, read_code_file("backend/app/repositories/base.py"), "python")

    add_section_heading(doc, "3.4. Бізнес-логіка у Workflow-класах")
    add_body_text(doc,
        "Бізнес-логіка винесена у окремі Workflow-класи, що забезпечує чітке розділення "
        "відповідальності. Кожен Workflow приймає вхідні дані через Pydantic-модель та "
        "повертає результат також у вигляді Pydantic-моделі. Залежності ін'єктуються "
        "через конструктор."
    )

    add_listing_caption(doc, "3.4", "Workflow створення групи (CreateGroupWorkflow)")
    add_code_annotation(doc,
        "CreateGroupWorkflow перевіряє ліміт груп для користувача, створює групу, "
        "додає творця як учасника з роллю Admin та призначає відповідні дозволи "
        "через GroupMemberPermissionRepository."
    )
    add_code_listing(doc, read_code_file("backend/app/workflows/group/create.py"), "python")

    add_body_text(doc,
        "Особливо важливим є Workflow життєвого циклу замовлення, який керує переходами "
        "між статусами та автоматично оновлює баланси при завершенні замовлення."
    )

    add_listing_caption(doc, "3.4", "Workflow життєвого циклу замовлення (фрагмент)")
    add_code_annotation(doc,
        "OrderLifecycleWorkflow.transition() перевіряє права (ініціатор або Editor), "
        "валідує перехід між статусами згідно VALID_TRANSITIONS, при завершенні "
        "автоматично розраховує та оновлює баланси учасників."
    )
    add_code_listing(doc, read_code_file("backend/app/workflows/order/lifecycle.py", end=92), "python")

    add_image(doc, order_path, "Діаграма станів життєвого циклу замовлення", width_cm=13)

    add_section_heading(doc, "3.5. API ендпоінти та маршрутизація")
    add_body_text(doc,
        "Ендпоінти організовано за допомогою APIRouter з чітким розділенням за ресурсами. "
        "Залежності (поточний користувач, репозиторії, workflow) ін'єктуються через Depends()."
    )

    add_listing_caption(doc, "3.5", "API ендпоінти управління групами (фрагмент)")
    add_code_annotation(doc,
        "Модуль groups.py визначає CRUD-ендпоінти для груп: list_groups (з урахуванням ролі), "
        "create_group (через CreateGroupWorkflow), get_group та інші. Залежності ін'єктуються "
        "через механізм Depends() фреймворку FastAPI."
    )
    add_code_listing(doc, read_code_file("backend/app/api/groups.py", end=67), "python")

    add_section_heading(doc, "3.6. Автентифікація та авторизація")
    add_body_text(doc,
        "Автентифікація реалізована через JWT-токени, що зберігаються в HTTP-only cookies "
        "для захисту від XSS-атак. Функція get_current_user є FastAPI-залежністю, що "
        "витягує токен з cookie, декодує його та повертає поточного користувача."
    )

    add_listing_caption(doc, "3.6", "Залежності автентифікації та авторизації")
    add_code_annotation(doc,
        "Функція get_current_user витягує access_token з cookie, декодує JWT через "
        "decode_access_token, знаходить користувача та перевіряє що акаунт активний. "
        "get_current_admin додатково перевіряє роль UserRole.ADMIN."
    )
    add_code_listing(doc,
        read_code_file("backend/app/dependencies.py", start=177, end=214),
        "python",
    )

    # ===========================================================
    #  РОЗДІЛ 4
    # ===========================================================
    add_chapter_heading(doc, "РОЗДІЛ 4. РОЗРОБКА КЛІЄНТСЬКОЇ ЧАСТИНИ")

    add_section_heading(doc, "4.1. Структура React-додатку")
    add_body_text(doc,
        "Клієнтська частина побудована на React 19 з TypeScript та Vite як збірником. "
        "Проєкт організовано за модульною структурою, де кожен функціональний модуль "
        "(auth, group, order, restaurant, balance, dashboard) містить свої сторінки "
        "та компоненти."
    )
    add_body_text(doc,
        "Стилізація здійснюється через Tailwind CSS v4 з бібліотекою компонентів shadcn/ui, "
        "яка надає доступні та стилізовані UI-компоненти (Button, Card, Dialog, Input, Label). "
        "Для іконок використовується Lucide React. Маршрутизація реалізована через React "
        "Router v7 з підтримкою захищених маршрутів (ProtectedRoute)."
    )

    add_section_heading(doc, "4.2. Управління станом з Redux Toolkit та RTK Query")
    add_body_text(doc,
        "Управління серверним станом реалізовано через RTK Query — інструмент Redux Toolkit "
        "для автоматичного кешування, інвалідації та повторного завантаження даних."
    )

    add_listing_caption(doc, "4.2", "Базова конфігурація RTK Query (baseApi.ts)")
    add_code_annotation(doc,
        "Модуль baseApi.ts створює базовий API-клієнт через createApi/fetchBaseQuery "
        "з автоматичною обробкою помилки 401 (перенаправлення на сторінку входу). "
        "Визначає теги для інвалідації кешу всіх ресурсів."
    )
    add_code_listing(doc, read_code_file("frontend/src/store/api/baseApi.ts"), "typescript")

    add_listing_caption(doc, "4.2", "API-сервіс замовлень (фрагмент orderApi.ts)")
    add_code_annotation(doc,
        "orderApi розширює базовий API через injectEndpoints(). Визначає query-ендпоінти "
        "(getOrders, getActiveOrder, getOrder) та мутації (createOrder, updateOrderStatus). "
        "providesTags/invalidatesTags забезпечують автоматичне оновлення кешу."
    )
    add_code_listing(doc, read_code_file("frontend/src/store/api/orderApi.ts", end=62), "typescript")

    add_section_heading(doc, "4.3. Сторінки та компоненти")
    add_body_text(doc,
        "Сторінки додатку реалізовано як React-компоненти, що використовують хуки RTK Query "
        "для завантаження даних. Кожна сторінка включає стани завантаження та помилки."
    )

    add_listing_caption(doc, "4.3", "Сторінка списку замовлень (фрагмент OrderListPage.tsx)")
    add_code_annotation(doc,
        "OrderListPage використовує useGetOrdersQuery та useGetRestaurantsQuery для "
        "отримання даних. Реалізує діалог створення замовлення з компонентом Combobox "
        "для вибору ресторану або введення нової назви."
    )
    add_code_listing(doc,
        read_code_file("frontend/src/modules/order/pages/OrderListPage.tsx", end=58),
        "typescript",
    )

    add_screenshot_placeholder(doc, "Головна сторінка користувача (дашборд)")
    add_screenshot_placeholder(doc, "Сторінка групи з аналітикою")
    add_screenshot_placeholder(doc, "Сторінка списку замовлень")
    add_screenshot_placeholder(doc, "Сторінка деталей замовлення")

    add_section_heading(doc, "4.4. Навігація та маршрутизація")
    add_body_text(doc,
        "Навігація реалізована через двопанельну бічну панель. Ліва вузька панель містить "
        "іконки: кнопку «Домашня сторінка», іконки груп та кнопку створення нової групи (+). "
        "Права панель є контекстно-залежною: при виборі домашньої сторінки відображаються "
        "пункти Home, Profile, Settings та Manage Users; при виборі групи — Dashboard, "
        "Members, Restaurants, Orders, Balances."
    )

    add_listing_caption(doc, "4.4", "Компонент бічної навігації (фрагмент Sidebar.tsx)")
    add_code_annotation(doc,
        "Sidebar визначає два контексти навігації: isHomeContext (домашні сторінки) "
        "та isGroupContext (сторінки групи). Використовує useGetGroupsQuery для "
        "відображення списку груп та useCreateGroupMutation для створення нових."
    )
    add_code_listing(doc,
        read_code_file("frontend/src/components/common/Layout/Sidebar.tsx", end=85),
        "typescript",
    )

    add_screenshot_placeholder(doc, "Бічна навігація — контекст домашньої сторінки")
    add_screenshot_placeholder(doc, "Бічна навігація — контекст групи")

    add_section_heading(doc, "4.5. UI-компоненти та дизайн-система shadcn/ui")
    add_body_text(doc,
        "Для побудови інтерфейсу використано бібліотеку shadcn/ui, яка надає набір доступних "
        "та стилізованих компонентів на базі Radix UI. Компоненти (Button, Card, Dialog, "
        "Input, Label, Popover) копіюються в проєкт та можуть бути кастомізовані. Додатково "
        "реалізовано власний компонент Combobox з можливістю пошуку та створення нових "
        "елементів, який використовується для вибору ресторану при створенні замовлення."
    )

    add_screenshot_placeholder(doc, "Компонент Combobox з можливістю створення")
    add_screenshot_placeholder(doc, "Діалог створення групи")

    # ===========================================================
    #  ВИСНОВКИ
    # ===========================================================
    add_chapter_heading(doc, "ВИСНОВКИ")

    add_body_text(doc,
        "За період проходження навчальної практики було особисто виконано повний цикл "
        "розробки веб-додатку «LunchTogether» — від аналізу вимог та проєктування архітектури "
        "до реалізації серверної та клієнтської частин. Додаток реалізує весь запланований "
        "функціонал: систему автентифікації з JWT, управління групами та учасниками, каталог "
        "ресторанів зі стравами, замовлення з повним життєвим циклом (Initiated → Confirmed → "
        "Ordered → Finished/Cancelled), систему балансів та аналітичні дашборди."
    )
    add_body_text(doc,
        "Серверна частина побудована на Python з використанням FastAPI, SQLAlchemy та PostgreSQL. "
        "Архітектурні рішення — шаблони Repository та Workflow — забезпечують чітке розділення "
        "відповідальності та спрощують подальший розвиток системи. Клієнтська частина реалізована "
        "на React з TypeScript та використовує Redux Toolkit з RTK Query для ефективного "
        "управління серверним станом. Інтерфейс побудовано на базі shadcn/ui та Tailwind CSS."
    )
    add_body_text(doc,
        "Реалізована дворівнева система прав та дозволів (рівень додатку та рівень групи) "
        "з п'ятьма типами групових дозволів та пресетами ролей забезпечує гнучке управління "
        "доступом у межах кожної групи."
    )
    add_body_text(doc,
        "Для подальшого вдосконалення системи пропонуються наступні напрямки розвитку:"
    )
    add_body_text(doc,
        "1) Впровадження сповіщень у реальному часі через WebSocket-з'єднання для "
        "миттєвого оповіщення учасників про зміни статусу замовлення та нові запрошення."
    )
    add_body_text(doc,
        "2) Інтеграція з платіжними системами (Monobank API, ПриватБанк API) для "
        "автоматизації розрахунків між учасниками без ручного коригування балансів."
    )
    add_body_text(doc,
        "3) Розробка мобільного додатку на базі React Native для забезпечення зручного "
        "доступу з мобільних пристроїв, що підвищить залученість користувачів."
    )
    add_body_text(doc,
        "4) Впровадження системи email-сповіщень для нагадування про активні замовлення "
        "та запрошення до груп."
    )
    add_body_text(doc,
        "5) Додавання підтримки кількох мов інтерфейсу (i18n) та темної теми для "
        "розширення аудиторії та покращення користувацького досвіду."
    )
    add_body_text(doc,
        "6) Впровадження кешування на рівні серверу (Redis) для підвищення продуктивності "
        "при збільшенні кількості активних користувачів."
    )
    add_body_text(doc,
        "Результати практики підтверджують доцільність обраних технологій та архітектурних "
        "рішень для побудови сучасних веб-додатків. Набуті практичні навички роботи з FastAPI, "
        "SQLAlchemy, React, Redux Toolkit та TypeScript є цінним досвідом для подальшої "
        "професійної діяльності."
    )

    # ===========================================================
    #  СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ
    # ===========================================================
    add_chapter_heading(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ")

    references = [
        "FastAPI Documentation. URL: https://fastapi.tiangolo.com/ (дата звернення: 10.02.2026).",
        "SQLAlchemy 2.0 Documentation. URL: https://docs.sqlalchemy.org/en/20/ (дата звернення: 10.02.2026).",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/ (дата звернення: 10.02.2026).",
        "Redux Toolkit Documentation. URL: https://redux-toolkit.js.org/ (дата звернення: 10.02.2026).",
        "shadcn/ui Documentation. URL: https://ui.shadcn.com/docs (дата звернення: 10.02.2026).",
        "React Documentation. URL: https://react.dev/ (дата звернення: 10.02.2026).",
        "Pydantic Documentation. URL: https://docs.pydantic.dev/latest/ (дата звернення: 10.02.2026).",
        "Alembic — Database Migration Tool. URL: https://alembic.sqlalchemy.org/en/latest/ (дата звернення: 10.02.2026).",
        "Tailwind CSS Documentation. URL: https://tailwindcss.com/docs (дата звернення: 10.02.2026).",
        "React Router Documentation. URL: https://reactrouter.com/ (дата звернення: 10.02.2026).",
    ]
    for i, ref in enumerate(references, start=1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = para.add_run(f"{i}. {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # ===========================================================
    #  ДОДАТКИ
    # ===========================================================
    add_chapter_heading(doc, "ДОДАТКИ")

    # ── ДОДАТОК А ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("ДОДАТОК А")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Структура файлів проєкту LunchTogether")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True

    project_tree = """\
LunchTogether/
├── backend/
│   ├── app/
│   │   ├── api/
│   │   │   ├── auth.py
│   │   │   ├── groups.py
│   │   │   ├── orders.py
│   │   │   ├── restaurants.py
│   │   │   ├── balances.py
│   │   │   ├── analytics.py
│   │   │   ├── users.py
│   │   │   └── router.py
│   │   ├── core/
│   │   │   ├── database.py
│   │   │   ├── security.py
│   │   │   ├── exceptions.py
│   │   │   └── middleware.py
│   │   ├── models/
│   │   │   ├── base.py
│   │   │   ├── enums.py
│   │   │   ├── user.py
│   │   │   ├── group.py
│   │   │   ├── order.py
│   │   │   ├── restaurant.py
│   │   │   └── balance.py
│   │   ├── repositories/
│   │   │   ├── base.py
│   │   │   ├── user.py
│   │   │   ├── group.py
│   │   │   ├── order.py
│   │   │   ├── restaurant.py
│   │   │   └── balance.py
│   │   ├── schemas/
│   │   │   ├── user.py
│   │   │   ├── group.py
│   │   │   ├── order.py
│   │   │   └── base.py
│   │   ├── workflows/
│   │   │   ├── group/
│   │   │   │   ├── create.py
│   │   │   │   ├── manage_members.py
│   │   │   │   └── invite.py
│   │   │   ├── order/
│   │   │   │   ├── create.py
│   │   │   │   └── lifecycle.py
│   │   │   └── balance/
│   │   │       └── adjust.py
│   │   ├── config.py
│   │   ├── dependencies.py
│   │   └── main.py
│   ├── alembic/
│   └── pyproject.toml
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── common/Layout/Sidebar.tsx
│   │   │   └── ui/ (Button, Card, Dialog, Input, ...)
│   │   ├── modules/
│   │   │   ├── auth/pages/
│   │   │   ├── group/pages/
│   │   │   ├── order/pages/
│   │   │   ├── restaurant/pages/
│   │   │   ├── balance/pages/
│   │   │   ├── dashboard/pages/
│   │   │   └── user/pages/
│   │   ├── store/
│   │   │   ├── api/ (baseApi, authApi, groupApi, ...)
│   │   │   └── slices/ (authSlice)
│   │   ├── hooks/
│   │   ├── types/
│   │   ├── constants/
│   │   └── App.tsx
│   ├── package.json
│   └── vite.config.ts
└── generate_report.py"""

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FAFAFA")
    _set_cell_borders(cell, color="CCCCCC")
    cell.paragraphs[0].clear()
    for i, line in enumerate(project_tree.strip().split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0)
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    # ===========================================================
    #  Save
    # ===========================================================
    doc.save(str(OUTPUT_FILE))
    print(f"Report saved to: {OUTPUT_FILE}")

    print("\n=== SCREENSHOT PLACEHOLDERS (insert images manually) ===")
    placeholders = [
        "Головна сторінка користувача (дашборд)",
        "Сторінка групи з аналітикою",
        "Сторінка списку замовлень",
        "Сторінка деталей замовлення",
        "Бічна навігація — контекст домашньої сторінки",
        "Бічна навігація — контекст групи",
        "Компонент Combobox з можливістю створення",
        "Діалог створення групи",
    ]
    for p in placeholders:
        print(f"  - Рис. {p}")

    print(f"\n=== GENERATED DIAGRAMS (embedded in report) ===")
    print(f"  - architecture.png -> {arch_path}")
    print(f"  - er_diagram.png -> {er_path}")
    print(f"  - permission_flow.png -> {perm_path}")
    print(f"  - order_lifecycle.png -> {order_path}")


if __name__ == "__main__":
    build_report()
